import io
import json
from pathlib import Path
import docx
import openpyxl
from pypdf import PageObject, PdfWriter
import pytest
import yaml

from career_radar.extractor import AnnouncementExtractor
from career_radar.fetcher import AnnouncementFetcher, AttachmentAccessError
from career_radar.models import (
    CandidateProfile,
    DimensionEvaluation,
    EvaluationResult,
    SourceObservation,
)
from career_radar.parser import AttachmentParser, HTMLAnnouncementParser
from career_radar.runner import finalize_evaluation_run


@pytest.fixture
def sample_xlsx_fixture(tmp_path: Path) -> Path:
    """Generates an in-memory synthetic XLSX attachment with 2 job positions."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "岗位表"

    # Header
    ws.append([
        "序号", "招聘单位", "工作部门", "岗位名称", "学历学位要求", "专业及代码",
        "年龄要求", "能力要求", "工作地点", "岗位等级"
    ])
    # Row 1 (Position 1)
    ws.append([
        "1",
        "广东药科大学",
        "信息工程学院",
        "数字媒体应用技术专任教师",
        "硕士研究生及以上学历并取得硕士学位",
        "计算机科学与技术（0812）、数字媒体（0854）",
        "35周岁以下",
        "能胜任UI交互设计、3D制作等专业课程",
        "广州",
        "专业技术岗位十一级以上",
    ])
    # Row 2 (Position 2)
    ws.append([
        "2",
        "广东药科大学",
        "数理学院",
        "理论物理学科带头人",
        "博士研究生并取得博士学位",
        "理论物理（070201）",
        "28周岁以下",
        "需具备海外全英文主讲经历",
        "广州",
        "专业技术岗位四级以上",
    ])

    xlsx_path = tmp_path / "attachment_posts.xlsx"
    wb.save(xlsx_path)
    return xlsx_path


@pytest.fixture
def sample_docx_fixture(tmp_path: Path) -> Path:
    """Generates a synthetic DOCX attachment with a job table."""
    doc = docx.Document()
    doc.add_heading("2026年公开招聘岗位明细表", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "岗位名称"
    hdr_cells[1].text = "招聘人数"
    hdr_cells[2].text = "学历要求"
    hdr_cells[3].text = "专业方向"
    hdr_cells[4].text = "年龄要求"
    hdr_cells[5].text = "能力要求"

    row_cells = table.add_row().cells
    row_cells[0].text = "人工智能与交互设计教师"
    row_cells[1].text = "1"
    row_cells[2].text = "硕士及以上"
    row_cells[3].text = "计算机科学/设计交叉方向"
    row_cells[4].text = "35周岁以下"
    row_cells[5].text = "具备跨学科教学经验"

    docx_path = tmp_path / "attachment_posts.docx"
    doc.save(docx_path)
    return docx_path


@pytest.fixture
def sample_pdf_fixture(tmp_path: Path) -> Path:
    """Generates a synthetic text-native PDF attachment."""
    writer = PdfWriter()
    page = PageObject.create_blank_page(width=612, height=792)
    writer.add_page(page)

    pdf_path = tmp_path / "attachment_notice.pdf"
    with open(pdf_path, "wb") as f:
        writer.write(f)
    return pdf_path


def test_attachment_parser_xlsx(sample_xlsx_fixture: Path):
    parser = AttachmentParser()
    parsed_tables = parser.parse_file(sample_xlsx_fixture)

    assert len(parsed_tables) >= 1
    table = parsed_tables[0]
    assert table["file_type"] == "xlsx"
    assert table["status"] == "success"
    assert len(table["rows"]) == 2
    assert table["rows"][0]["cells"]["岗位名称"] == "数字媒体应用技术专任教师"
    assert "35周岁以下" in table["rows"][0]["cells"]["年龄要求"]
    assert table["rows"][1]["cells"]["岗位名称"] == "理论物理学科带头人"


def test_attachment_parser_docx(sample_docx_fixture: Path):
    parser = AttachmentParser()
    parsed_tables = parser.parse_file(sample_docx_fixture)

    assert len(parsed_tables) >= 1
    table = parsed_tables[0]
    assert table["file_type"] == "docx"
    assert table["status"] == "success"
    assert len(table["rows"]) == 1
    assert table["rows"][0]["cells"]["岗位名称"] == "人工智能与交互设计教师"


def test_attachment_parser_pdf(sample_pdf_fixture: Path):
    parser = AttachmentParser()
    parsed_tables = parser.parse_file(sample_pdf_fixture)

    assert len(parsed_tables) >= 1
    table = parsed_tables[0]
    assert table["file_type"] == "pdf"
    assert table["status"] == "success"
    assert "pages" in table


def test_attachment_parser_legacy_unsupported(tmp_path: Path):
    """Verifies that legacy .xls and .doc files return unsupported_legacy_format without crashing."""
    xls_path = tmp_path / "legacy.xls"
    xls_path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")  # mock legacy OLE header

    doc_path = tmp_path / "legacy.doc"
    doc_path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")

    parser = AttachmentParser()
    res_xls = parser.parse_file(xls_path)
    assert len(res_xls) == 1
    assert res_xls[0]["status"] == "unsupported_legacy_format"
    assert res_xls[0]["rows"] == []

    res_doc = parser.parse_file(doc_path)
    assert len(res_doc) == 1
    assert res_doc[0]["status"] == "unsupported_legacy_format"
    assert res_doc[0]["rows"] == []


def test_announcement_extractor_slices_one_announcement_to_n_observations(
    tmp_path: Path, sample_xlsx_fixture: Path
):
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head><title>广东药科大学2026年人才招聘公告</title></head>
    <body>
      <div class="article-title">广东药科大学2026年公开招聘专任教师公告</div>
      <div class="publish-time">2026-08-15 09:00</div>
      <div class="content">
        <p>为满足教学科研需要，我校现面向社会公开招聘专任教师。</p>
        <div class="attachment">
          <a href="{sample_xlsx_fixture.as_uri()}">附件1：岗位需求表.xlsx</a>
        </div>
      </div>
    </body>
    </html>
    """

    extractor = AnnouncementExtractor(cache_dir=tmp_path / ".data" / "announcements")
    observations = extractor.extract_from_html_and_attachments(
        html_content=html_content,
        source_url="http://hrss.gd.gov.cn/zwgk/sydwzp/zpgg/content_202601.html",
        source_id="gd_hrss_official",
        source_name="广东省人力资源和社会保障厅",
        local_attachment_paths=[sample_xlsx_fixture],
        recruiting_organization="广东药科大学",
    )

    # 1 Announcement -> 2 SourceObservations
    assert len(observations) == 2

    obs_1 = observations[0]
    assert obs_1.job_title == "数字媒体应用技术专任教师"
    # Organization must be canonical institution, not department
    assert obs_1.organization == "广东药科大学"
    assert obs_1.location == "广州"
    # Track must NOT be mapped from 岗位等级
    assert obs_1.track == ""
    assert "35周岁以下" in obs_1.extracted_requirements["age_text"]
    assert "硕士研究生" in obs_1.extracted_requirements["education_text"]
    assert obs_1.provenance is not None
    assert obs_1.provenance["file_name"] == "attachment_posts.xlsx"
    assert obs_1.provenance["row_index"] == 2
    assert obs_1.provenance["department"] == "信息工程学院"
    assert obs_1.provenance["raw_cells"]["岗位等级"] == "专业技术岗位十一级以上"

    obs_2 = observations[1]
    assert obs_2.job_title == "理论物理学科带头人"
    assert obs_2.organization == "广东药科大学"
    assert obs_2.provenance["department"] == "数理学院"
    assert "28周岁以下" in obs_2.extracted_requirements["age_text"]
    assert "博士研究生" in obs_2.extracted_requirements["education_text"]


def test_announcement_extractor_no_silent_fallback_on_failed_attachments(tmp_path: Path):
    """
    CRITICAL BLOCKER 2 TEST:
    If an announcement references attachments but no valid job table rows can be extracted,
    it MUST NOT create fake observations from the announcement title. It must return [].
    """
    html_content = """
    <html>
      <head><title>广东某大学2026年公开招聘工作人员公告</title></head>
      <body>
        <h1>广东某大学2026年公开招聘工作人员公告</h1>
        <p>具体招聘岗位详见附件。</p>
      </body>
    </html>
    """

    extractor = AnnouncementExtractor(cache_dir=tmp_path / ".data" / "announcements")
    observations = extractor.extract_from_html_and_attachments(
        html_content=html_content,
        source_url="https://rsc.example.edu.cn/recruit/001.html",
        source_id="example_rsc",
        source_name="某大学人事处",
        local_attachment_paths=[],
    )

    # STRICT ASSERTION: Zero fake observations created!
    assert observations == []


def test_first_party_announcement_to_daily_digest_seam(
    tmp_path: Path, sample_xlsx_fixture: Path
):
    """
    Highest Seam Verification:
    HTML Announcement + Attachment XLSX
    -> Extractor slices into 2 SourceObservations
    -> Agent decisions
    -> finalize_evaluation_run
    -> .data/opportunities.jsonl & reports/YYYY-MM-DD.md
    """
    html_content = f"""
    <html>
      <body>
        <h1>广东药科大学2026年招聘</h1>
        <a href="{sample_xlsx_fixture.as_uri()}">附件.xlsx</a>
      </body>
    </html>
    """
    extractor = AnnouncementExtractor(cache_dir=tmp_path / ".data" / "announcements")
    observations = extractor.extract_from_html_and_attachments(
        html_content=html_content,
        source_url="http://hrss.gd.gov.cn/zwgk/sydwzp/zpgg/content_test.html",
        source_id="gd_hrss_official",
        source_name="广东省人力资源和社会保障厅",
        local_attachment_paths=[sample_xlsx_fixture],
        recruiting_organization="广东药科大学",
    )
    assert len(observations) == 2

    # Synthetic Agent decisions for the 2 extracted observations
    now = "2026-08-15T10:00:00+08:00"
    decisions = [
        EvaluationResult(
            final_recommendation="建议关注",
            dimension_evaluations={
                "Age": DimensionEvaluation("Age", "PASS", observations[0].extracted_requirements["age_text"], "符合年龄"),
                "Education": DimensionEvaluation("Education", "PASS", observations[0].extracted_requirements["education_text"], "符合学历"),
                "Formal Qualification": DimensionEvaluation("Formal Qualification", "PASS", observations[0].extracted_requirements["formal_qualification_text"], "契合专业"),
                "Capability Fit": DimensionEvaluation("Capability Fit", "PASS", observations[0].extracted_requirements["capability_fit_text"], "具备胜任力"),
                "Teaching Experience": DimensionEvaluation("Teaching Experience", "PASS", "", "具备带教经历"),
                "Industry Experience": DimensionEvaluation("Industry Experience", "PASS", "", "具备行业背景"),
            },
            evaluated_at=now,
        ),
        EvaluationResult(
            final_recommendation="明显不符合",
            dimension_evaluations={
                "Age": DimensionEvaluation("Age", "FAIL", observations[1].extracted_requirements["age_text"], "年龄超限"),
                "Education": DimensionEvaluation("Education", "FAIL", observations[1].extracted_requirements["education_text"], "博士要求不满足"),
                "Formal Qualification": DimensionEvaluation("Formal Qualification", "FAIL", observations[1].extracted_requirements["formal_qualification_text"], "专业不符"),
                "Capability Fit": DimensionEvaluation("Capability Fit", "FAIL", "", "不匹配"),
                "Teaching Experience": DimensionEvaluation("Teaching Experience", "FAIL", "", "无海外经验"),
                "Industry Experience": DimensionEvaluation("Industry Experience", "N/A", "", "不适用"),
            },
            evaluated_at=now,
        ),
    ]

    summary = finalize_evaluation_run(
        observations=observations,
        evaluation_results=decisions,
        data_dir=tmp_path / ".data",
        reports_dir=tmp_path / "reports",
        run_date="2026-08-15",
    )

    assert summary["success"] is True
    assert summary["total_evaluated"] == 2
    assert summary["recommended_count"] == 1
    assert summary["mismatch_count"] == 1

    # Check persistence
    opps_file = tmp_path / ".data" / "opportunities.jsonl"
    assert opps_file.exists()
    with open(opps_file, "r", encoding="utf-8") as f:
        records = [json.loads(line) for line in f if line.strip()]
    assert len(records) == 2
    assert records[0]["job_title"] == "数字媒体应用技术专任教师"
    assert records[0]["organization"] == "广东药科大学"
    assert records[0]["latest_evaluation"]["final_recommendation"] == "建议关注"

    # Check report
    report_file = tmp_path / "reports" / "2026-08-15.md"
    assert report_file.exists()
    content = report_file.read_text(encoding="utf-8")
    assert "数字媒体应用技术专任教师" in content
    assert "广东药科大学" in content
    assert "http://hrss.gd.gov.cn/zwgk/sydwzp/zpgg/content_test.html" in content
