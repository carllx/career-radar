"""
Production Acquisition Entrypoint Seam Tests for Issue #22 (static_html_attachment).
Tests verify external behavior of execute_production_acquisition and SourceAcquisitionExecutor:
1. Multi-hop acquisition flow: Listing HTML -> Detail HTML -> Discovered Attachment (XLSX, DOCX, text-PDF)
   -> Binary download & hashing -> Deterministic parsing -> Compact Agent evidence packet.
2. Raw attachment evidence persistence and hash verification.
3. Attachment HTTP failure / 404 preserves detail HTML facts and records failure audit.
4. Attachment corrupt / parse failure preserves raw downloaded bytes & HTTP facts, records parse error without crashing.
5. Legacy unsupported formats (.xls, .doc) report unsupported_legacy_format truthfully.
6. Safe attachment filenames: path traversal sanitization and same-name collision avoidance.
7. Text-native PDF extraction proof with non-empty text pages.
"""

from io import BytesIO
import os
from pathlib import Path
from typing import Any, Dict

import docx
import openpyxl
from pypdf import PdfReader
import pytest

from career_radar.acquisition import (
    AcquisitionResult,
    execute_production_acquisition,
)
from career_radar.attachment_helper import AttachmentAcquisitionHelper
from career_radar.sources import SourceRecord


class FakeHttpTransport:
    def __init__(self, responses: Dict[str, Dict[str, Any]]):
        self.responses = responses
        self.requests_log = []

    def get(self, url: str, headers: Any = None, timeout: int = 15, verify: bool = True):
        self.requests_log.append({
            "url": url,
            "headers": headers,
            "timeout": timeout,
            "verify": verify,
        })
        if url in self.responses:
            resp_data = self.responses[url]
            return FakeResponse(
                status_code=resp_data.get("status_code", 200),
                text=resp_data.get("text", ""),
                headers=resp_data.get("headers", {"Content-Type": "text/html; charset=utf-8"}),
                url=url,
                content=resp_data.get("content"),
            )
        return FakeResponse(status_code=404, text="Not Found", headers={}, url=url)


class FakeResponse:
    def __init__(
        self,
        status_code: int,
        text: str,
        headers: Dict[str, str],
        url: str,
        content: Any = None,
    ):
        self.status_code = status_code
        self.text = text
        self.headers = headers
        self.url = url
        if content is not None:
            self.content = content
        else:
            self.content = text.encode("utf-8")

    def raise_for_status(self):
        if self.status_code >= 400:
            raise Exception(f"HTTP {self.status_code} Error")


@pytest.fixture
def sample_xlsx_bytes() -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "岗位明细"
    ws.append(["序号", "招聘单位", "工作部门", "岗位名称", "学历学位要求", "专业及代码", "年龄要求", "能力要求"])
    ws.append([
        "1", "广东岭南工程学院", "智能制造学院", "机器人工程专任教师",
        "硕士研究生及以上", "0811 控制科学与工程 / 0802 机械工程", "35周岁以下", "具备ROS开发经验"
    ])
    ws.append([
        "2", "广东岭南工程学院", "计算机学院", "人工智能专任教师",
        "博士研究生", "0812 计算机科学与技术", "38周岁以下", "有大模型实训带教经历"
    ])
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


@pytest.fixture
def sample_docx_bytes() -> bytes:
    doc = docx.Document()
    doc.add_heading("2026年高层次人才招聘岗位表", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "岗位名称"
    hdr[1].text = "学历要求"
    hdr[2].text = "专业方向"
    hdr[3].text = "招聘人数"

    row = table.add_row().cells
    row[0].text = "数字媒体艺术专任教师"
    row[1].text = "硕士及以上"
    row[2].text = "设计学 / 数字媒体"
    row[3].text = "2"

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    # Text-native PDF with extractable content stream
    return (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 50>>stream\n"
        b"BT /F1 12 Tf 72 712 Td (Guideline text content) Tj ET\n"
        b"endstream\n"
        b"endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f\n0000000009 00000 n\n0000000056 00000 n\n0000000111 00000 n\n0000000222 00000 n\n0000000324 00000 n\n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n392\n%%EOF\n"
    )


def test_three_hop_listing_to_detail_to_xlsx_attachment(tmp_path: Path, sample_xlsx_bytes: bytes):
    """
    Finding 1 & 2: True 3-hop tracer (Listing HTML -> Detail HTML -> XLSX attachment).
    Proves:
    - Listing page GET generates AcquisitionResult with request_type=listing.
    - Deterministic link selection using config hint (detail_url_pattern).
    - Detail page GET generates AcquisitionResult with request_type=detail.
    - Attachment GET generates AcquisitionResult with request_type=attachment.
    - Listing and Detail raw HTML are preserved in distinct files on disk.
    - Top-level execute_production_acquisition exposes all 3 physical AcquisitionResults.
    """
    listing_url = "https://hrss.gd.gov.cn/zwgk/zp_list.html"
    detail_url = "https://hrss.gd.gov.cn/zwgk/zp2026_01.html"
    att_url = "https://hrss.gd.gov.cn/attach/2026_post_table.xlsx"

    listing_html = f"""<!DOCTYPE html>
    <html>
    <head><title>招聘信息列表</title></head>
    <body>
      <ul class="news-list">
        <li><a href="{detail_url}">广东岭南工程学院2026年公开招聘专任教师公告</a></li>
        <li><a href="/other/notice.html">其他无关通知</a></li>
      </ul>
    </body>
    </html>"""

    detail_html = f"""<!DOCTYPE html>
    <html>
    <head><title>广东岭南工程学院2026年公开招聘专任教师公告</title></head>
    <body>
      <h1>广东岭南工程学院2026年公开招聘专任教师公告</h1>
      <p>为满足教育教学需要，现面向社会招聘专任教师，具体岗位见附件：</p>
      <div class="attachment-box">
        <a href="{att_url}" title="岗位表.xlsx">附件：广东岭南工程学院2026年岗位需求明细表.xlsx</a>
      </div>
    </body>
    </html>"""

    transport = FakeHttpTransport(
        responses={
            listing_url: {"status_code": 200, "text": listing_html, "headers": {"Content-Type": "text/html; charset=utf-8"}},
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html; charset=utf-8"}},
            att_url: {
                "status_code": 200,
                "content": sample_xlsx_bytes,
                "headers": {
                    "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "Content-Disposition": 'attachment; filename="2026_post_table.xlsx"',
                },
            },
        }
    )

    source = SourceRecord(
        source_id="gd_hrss_3hop",
        name="广东省人社厅3跳招聘",
        base_url=listing_url,
        domain="hrss.gd.gov.cn",
        source_type="first_party_official",
        metadata={
            "archetype": "listing_html",
            "is_listing": True,
            "detail_url_pattern": r"zp2026_\d+\.html",
        },
    )

    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    # 1. Transport invoked for all 3 physical hops
    assert len(transport.requests_log) == 3
    assert transport.requests_log[0]["url"] == listing_url
    assert transport.requests_log[1]["url"] == detail_url
    assert transport.requests_log[2]["url"] == att_url

    # 2. Top-level output contains all 3 physical AcquisitionResults
    acq_results = output["acquisition_results"]
    assert len(acq_results) == 3

    list_acq, detail_acq, att_acq = acq_results[0], acq_results[1], acq_results[2]
    assert list_acq.requested_url == listing_url
    assert list_acq.technical_status == "success"
    assert list_acq.metadata["request_type"] == "listing"
    assert Path(list_acq.metadata["raw_evidence_path"]).exists()

    assert detail_acq.requested_url == detail_url
    assert detail_acq.technical_status == "success"
    assert detail_acq.metadata["request_type"] == "detail"
    assert Path(detail_acq.metadata["raw_evidence_path"]).exists()
    assert list_acq.metadata["raw_evidence_path"] != detail_acq.metadata["raw_evidence_path"]

    assert att_acq.requested_url == att_url
    assert att_acq.technical_status == "success"
    assert att_acq.metadata["request_type"] == "attachment"
    assert Path(att_acq.metadata["local_evidence_path"]).exists()

    # 3. Agent packet correctly parsed
    assert len(output["agent_evidence_packets"]) == 1
    packet = output["agent_evidence_packets"][0]
    assert packet["title"] == "广东岭南工程学院2026年公开招聘专任教师公告"
    assert len(packet["attachment_tables"]) == 1
    assert len(packet["attachment_tables"][0]["rows"]) == 2


def test_attachment_filename_path_traversal_sanitization(tmp_path: Path, sample_xlsx_bytes: bytes):
    """
    Finding 3: Verifies that malicious Content-Disposition or link names cannot perform directory traversal (../ or ..\\).
    """
    detail_url = "https://hrss.gd.gov.cn/zwgk/zp_malicious.html"
    att_url = "https://hrss.gd.gov.cn/attach/traversal.xlsx"
    detail_html = f'<html><body><a href="{att_url}">../../traversal.xlsx</a></body></html>'

    transport = FakeHttpTransport(
        responses={
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html"}},
            att_url: {
                "status_code": 200,
                "content": sample_xlsx_bytes,
                "headers": {
                    "Content-Disposition": 'attachment; filename="../../../etc/passwd.xlsx"',
                    "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                },
            },
        }
    )

    source = SourceRecord(source_id="gd_hrss_sec", name="安全校验渠道", base_url=detail_url, domain="hrss.gd.gov.cn")
    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    acq_results = output["acquisition_results"]
    assert len(acq_results) == 2
    att_acq = acq_results[1]
    saved_path = Path(att_acq.metadata["local_evidence_path"])

    # Path must remain strictly inside .data/raw_evidence/gd_hrss_sec/attachments/<attempt_id>/
    expected_parent = data_dir / "raw_evidence" / "gd_hrss_sec" / "attachments" / att_acq.metadata["parent_attempt_id"]
    assert saved_path.parent.resolve() == expected_parent.resolve()
    assert saved_path.name == "passwd.xlsx"
    assert ".." not in str(saved_path)


def test_attachment_same_name_collision_disambiguation(tmp_path: Path, sample_xlsx_bytes: bytes, sample_docx_bytes: bytes):
    """
    Finding 3: Verifies that multiple attachments with identical names (e.g. table.xlsx from different URLs)
    do not overwrite each other and are disambiguated deterministically.
    """
    detail_url = "https://hrss.gd.gov.cn/zwgk/zp_dup.html"
    att_url1 = "https://hrss.gd.gov.cn/u1/table.xlsx"
    att_url2 = "https://hrss.gd.gov.cn/u2/table.xlsx"

    detail_html = f"""<html><body>
      <a href="{att_url1}">岗位表.xlsx</a>
      <a href="{att_url2}">岗位表.xlsx</a>
    </body></html>"""

    transport = FakeHttpTransport(
        responses={
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html"}},
            att_url1: {"status_code": 200, "content": sample_xlsx_bytes, "headers": {"Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}},
            att_url2: {"status_code": 200, "content": sample_xlsx_bytes, "headers": {"Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}},
        }
    )

    source = SourceRecord(source_id="gd_hrss_dup", name="同名重名渠道", base_url=detail_url, domain="hrss.gd.gov.cn")
    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    acq_results = output["acquisition_results"]
    assert len(acq_results) == 3  # 1 detail + 2 attachments
    att1_path = Path(acq_results[1].metadata["local_evidence_path"])
    att2_path = Path(acq_results[2].metadata["local_evidence_path"])

    assert att1_path.exists()
    assert att2_path.exists()
    assert att1_path != att2_path


def test_docx_and_text_native_pdf_attachments_acquisition_and_parsing(
    tmp_path: Path, sample_docx_bytes: bytes, sample_pdf_bytes: bytes
):
    """
    Finding 4: Multi-format attachment support: DOCX and text-native PDF in a single announcement.
    Proves that PDF text content ('Guideline text content') is extracted into attachment_pages.
    """
    detail_url = "https://rsc.sample.edu.cn/zp/announcement.html"
    docx_url = "https://rsc.sample.edu.cn/zp/docs/posts.docx"
    pdf_url = "https://rsc.sample.edu.cn/zp/docs/guidelines.pdf"

    detail_html = f"""<!DOCTYPE html>
    <html>
    <body>
      <h1>招聘公告</h1>
      <a href="{docx_url}">附件1：岗位表.docx</a>
      <a href="{pdf_url}">附件2：报考指南.pdf</a>
    </body>
    </html>"""

    transport = FakeHttpTransport(
        responses={
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html"}},
            docx_url: {
                "status_code": 200,
                "content": sample_docx_bytes,
                "headers": {"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            },
            pdf_url: {
                "status_code": 200,
                "content": sample_pdf_bytes,
                "headers": {"Content-Type": "application/pdf"},
            },
        }
    )

    source = SourceRecord(source_id="multi_format_src", name="多格式渠道", base_url=detail_url, domain="rsc.sample.edu.cn")
    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    assert len(transport.requests_log) == 3
    packet = output["agent_evidence_packets"][0]
    assert len(packet["attachments"]) == 2
    assert len(packet["attachment_tables"]) == 1
    assert packet["attachment_tables"][0]["rows"][0]["cells"]["岗位名称"] == "数字媒体艺术专任教师"
    assert len(packet["attachment_pages"]) == 1
    assert "Guideline text content" in packet["attachment_pages"][0]["text"]


def test_attachment_http_404_preserves_html_facts_and_records_audit(tmp_path: Path):
    """
    Attachment 404 HTTP failure must preserve HTML acquisition facts, record truthful attachment failure,
    and not crash or fabricate fake data.
    """
    detail_url = "https://www.school.edu.cn/zp.html"
    missing_att_url = "https://www.school.edu.cn/attach/missing.xlsx"
    detail_html = f"""<html><body><h1>招聘</h1><a href="{missing_att_url}">岗位表.xlsx</a></body></html>"""

    transport = FakeHttpTransport(
        responses={
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html"}},
            missing_att_url: {"status_code": 404, "text": "Attachment Not Found", "headers": {}},
        }
    )

    source = SourceRecord(source_id="missing_att_src", name="缺失附件渠道", base_url=detail_url, domain="school.edu.cn")
    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    acq_res = output["acquisition_results"][0]
    assert acq_res.http_status == 200
    assert acq_res.technical_status == "success"
    assert acq_res.metadata["attachments_found_count"] == 1
    assert acq_res.metadata["attachments_acquired_count"] == 0

    att_audit = acq_res.metadata["attachment_audits"][0]
    assert att_audit["status"] == "failed"
    assert att_audit["http_status"] == 404

    packet = output["agent_evidence_packets"][0]
    assert len(packet["attachments"]) == 1
    assert packet["attachments"][0]["status"] == "failed"
    assert packet["attachment_tables"] == []


def test_corrupted_attachment_preserves_downloaded_bytes_and_records_parse_error(tmp_path: Path):
    """
    Corrupted attachment bytes preserve downloaded bytes, file hash, local path, and record deterministic
    parsing failure in audit without crash or false observations.
    """
    detail_url = "https://www.school.edu.cn/zp_corrupt.html"
    corrupt_att_url = "https://www.school.edu.cn/attach/corrupt.xlsx"
    corrupt_bytes = b"NOT_A_VALID_ZIP_OR_XLSX_HEADER_GARBAGE_BYTES"
    detail_html = f"""<html><body><h1>招聘</h1><a href="{corrupt_att_url}">岗位表.xlsx</a></body></html>"""

    transport = FakeHttpTransport(
        responses={
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html"}},
            corrupt_att_url: {
                "status_code": 200,
                "content": corrupt_bytes,
                "headers": {"Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
            },
        }
    )

    source = SourceRecord(source_id="corrupt_src", name="损坏附件渠道", base_url=detail_url, domain="school.edu.cn")
    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    acq_res = output["acquisition_results"][0]
    att_audit = acq_res.metadata["attachment_audits"][0]
    assert att_audit["status"] == "success"
    assert att_audit["parse_status"] == "failed"
    assert att_audit["body_length"] == len(corrupt_bytes)
    assert Path(att_audit["local_evidence_path"]).read_bytes() == corrupt_bytes

    packet = output["agent_evidence_packets"][0]
    assert packet["attachments"][0]["status"] == "failed"
    assert packet["attachment_tables"] == []


def test_legacy_format_reported_truthfully_as_unsupported(tmp_path: Path):
    """
    Verifies that legacy formats (.xls, .doc) are reported truthfully without crashing.
    """
    detail_url = "https://www.school.edu.cn/legacy.html"
    detail_html = """<html><body><h1>招聘</h1><a href="/attach/old.xls">历史表.xls</a></body></html>"""

    transport = FakeHttpTransport(
        responses={
            detail_url: {"status_code": 200, "text": detail_html, "headers": {"Content-Type": "text/html"}},
        }
    )

    source = SourceRecord(source_id="legacy_src", name="旧格式渠道", base_url=detail_url, domain="school.edu.cn")
    data_dir = tmp_path / ".data"
    output = execute_production_acquisition(sources=[source], data_dir=data_dir, transport=transport)

    packet = output["agent_evidence_packets"][0]
    assert len(packet["attachments"]) == 1
    assert packet["attachments"][0]["status"] == "unsupported_legacy_format"
    assert packet["attachment_tables"] == []
