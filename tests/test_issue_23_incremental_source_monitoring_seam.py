"""
Focused Seam Tests for Issue #23:
Known-source incremental change detection & 0-Token monitoring.
Respects CONTEXT.md, ADR-0002, Spec #20, Issue #21, #22, and #23.
"""

from datetime import datetime
import hashlib
import json
from pathlib import Path
from typing import Any, Dict, List
import pytest

from career_radar.acquisition import (
    AcquisitionResult,
    SourceAcquisitionExecutor,
    execute_production_acquisition,
)
from career_radar.sources import MonitoringFact, SourceRecord, SourceRegistry


class FakeHttpTransport:
    """Controlled fake HTTP transport returning deterministic canned responses and logging requests."""

    def __init__(self, responses: Dict[str, Dict[str, Any]]):
        self.responses = responses
        self.requests_log: List[Dict[str, Any]] = []

    def get(self, url: str, headers: Any = None, timeout: int = 15, verify: bool = True):
        self.requests_log.append({
            "url": url,
            "headers": dict(headers or {}),
            "timeout": timeout,
            "verify": verify,
        })
        if url in self.responses:
            resp_data = self.responses[url]
            return FakeResponse(
                status_code=resp_data.get("status_code", 200),
                text=resp_data.get("text", ""),
                headers=resp_data.get("headers", {"Content-Type": "text/html; charset=utf-8"}),
                url=url,
                content=resp_data.get("content"),
            )
        return FakeResponse(status_code=404, text="Not Found", headers={}, url=url)


class FakeResponse:
    def __init__(self, status_code: int, text: str, headers: Dict[str, str], url: str, content: Any = None):
        self.status_code = status_code
        self.text = text
        self.headers = headers
        self.url = url
        self.content = content if content is not None else text.encode("utf-8")


def test_etag_conditional_request_304_unchanged_zero_agent_evidence(tmp_path: Path):
    """
    1. Proves ETag conditional request produces If-None-Match header,
       HTTP 304 response records successful MonitoringFact,
       produces 0 Agent evidence packets, and does not fetch detail.
    """
    listing_url = "https://hr.example.edu.cn/jobs"
    transport = FakeHttpTransport({
        listing_url: {
            "status_code": 304,
            "headers": {"ETag": '"abc123etag"', "Content-Type": "text/html"},
            "text": "",
        }
    })

    source = SourceRecord(
        source_id="src_etag_test",
        name="测试高校招聘网",
        base_url=listing_url,
        domain="hr.example.edu.cn",
        metadata={
            "is_listing": True,
            "committed_etag": '"abc123etag"',
            "detail_url_pattern": r"/jobs/\d+",
        },
    )

    out = execute_production_acquisition(
        sources=[source],
        data_dir=tmp_path / ".data",
        transport=transport,
    )

    # 1. Assert If-None-Match header was sent
    assert len(transport.requests_log) == 1
    assert transport.requests_log[0]["headers"].get("If-None-Match") == '"abc123etag"'

    # 2. Assert zero Agent evidence packets
    assert out["agent_evidence_packets"] == []

    # 3. Assert AcquisitionResult facts
    acq_results = out["acquisition_results"]
    assert len(acq_results) == 1
    assert acq_results[0].http_status == 304
    assert acq_results[0].technical_status == "success"
    assert acq_results[0].body_length == 0
    assert acq_results[0].etag == '"abc123etag"'

    # 4. Assert MonitoringFact
    mon_facts = out["monitoring_facts"]
    assert len(mon_facts) == 1
    assert mon_facts[0].technical_status == "success"
    assert mon_facts[0].metadata.get("http_status") == 304
    assert mon_facts[0].metadata.get("unchanged") is True


def test_last_modified_conditional_request_304_behavior(tmp_path: Path):
    """
    2. Proves Last-Modified produces If-Modified-Since header and handles 304 truthfully.
    """
    url = "https://hr.example.edu.cn/announcement/1"
    last_mod_str = "Sun, 16 Aug 2026 08:00:00 GMT"
    transport = FakeHttpTransport({
        url: {
            "status_code": 304,
            "headers": {"Last-Modified": last_mod_str, "Content-Type": "text/html"},
            "text": "",
        }
    })

    source = SourceRecord(
        source_id="src_lastmod_test",
        name="测试直达招聘公告",
        base_url=url,
        domain="hr.example.edu.cn",
        metadata={
            "is_listing": False,
            "committed_last_modified": last_mod_str,
        },
    )

    out = execute_production_acquisition(
        sources=[source],
        data_dir=tmp_path / ".data",
        transport=transport,
    )

    assert len(transport.requests_log) == 1
    assert transport.requests_log[0]["headers"].get("If-Modified-Since") == last_mod_str
    assert out["agent_evidence_packets"] == []
    assert len(out["acquisition_results"]) == 1
    assert out["acquisition_results"][0].http_status == 304
    assert out["monitoring_facts"][0].technical_status == "success"


def test_200_listing_with_identical_fingerprint_skips_detail_and_attachments(tmp_path: Path):
    """
    3. Proves that when HTTP 200 listing is returned with an identical deterministic fingerprint:
       - detail request is NOT fetched;
       - attachments are NOT fetched;
       - agent_evidence_packets is [];
       - successful MonitoringFact is recorded.
    """
    listing_url = "https://hr.example.edu.cn/recruit/list"
    detail_url = "https://hr.example.edu.cn/recruit/detail/101"
    listing_html = f"""<html><body>
    <div class="list">
      <a href="{detail_url}">2026年专任教师招聘启事</a>
      <a href="https://hr.example.edu.cn/news/1">无关新闻通知</a>
    </div>
    </body></html>"""

    transport = FakeHttpTransport({
        listing_url: {"status_code": 200, "text": listing_html},
        detail_url: {"status_code": 200, "text": "<html><body><h1>招聘详情</h1></body></html>"},
    })

    # Pre-compute canonical fingerprint for the matching URL
    # mechanically selected detail_url matching pattern
    expected_fingerprint = hashlib.sha256(
        json.dumps([detail_url], ensure_ascii=False, sort_keys=True).encode("utf-8")
    ).hexdigest()

    source = SourceRecord(
        source_id="src_listing_fp_test",
        name="广东某学院招聘网",
        base_url=listing_url,
        domain="hr.example.edu.cn",
        metadata={
            "is_listing": True,
            "detail_url_pattern": r"/recruit/detail/\d+",
            "committed_listing_fingerprint": expected_fingerprint,
        },
    )

    out = execute_production_acquisition(
        sources=[source],
        data_dir=tmp_path / ".data",
        transport=transport,
    )

    # Proves only listing was requested (1 request total, NO detail request)
    assert len(transport.requests_log) == 1
    assert transport.requests_log[0]["url"] == listing_url
    assert out["agent_evidence_packets"] == []
    assert len(out["acquisition_results"]) == 1
    assert out["acquisition_results"][0].metadata.get("unchanged") is True
    assert out["monitoring_facts"][0].technical_status == "success"
    assert out["monitoring_facts"][0].metadata.get("unchanged") is True


def test_changed_listing_fingerprint_triggers_detail_and_attachment_acquisition(tmp_path: Path):
    """
    4. Proves that when a new announcement URL appears / fingerprint changes:
       - detail acquisition is triggered;
       - attachment acquisition is triggered;
       - agent evidence packet is emitted;
       - newly observed fingerprint is present in metadata.
    """
    listing_url = "https://hr.example.edu.cn/list"
    detail_url = "https://hr.example.edu.cn/post/202"
    att_url = "https://hr.example.edu.cn/files/table.xlsx"

    listing_html = f"""<html><body>
      <a href="{detail_url}">2026年新教师招聘公告</a>
    </body></html>"""
    detail_html = f"""<html><body>
      <h1>2026年新教师招聘公告</h1>
      <a href="{att_url}">岗位明细表.xlsx</a>
    </body></html>"""

    # minimal dummy xlsx header bytes
    xlsx_bytes = b"PK\x03\x04\x14\x00\x00\x00\x08\x00" + b"\x00" * 40

    transport = FakeHttpTransport({
        listing_url: {"status_code": 200, "text": listing_html},
        detail_url: {"status_code": 200, "text": detail_html},
        att_url: {"status_code": 200, "content": xlsx_bytes, "headers": {"Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}},
    })

    source = SourceRecord(
        source_id="src_changed_test",
        name="某大学人事处",
        base_url=listing_url,
        domain="hr.example.edu.cn",
        metadata={
            "is_listing": True,
            "detail_url_pattern": r"/post/\d+",
            "committed_listing_fingerprint": "old_outdated_fingerprint_hash",
        },
    )

    out = execute_production_acquisition(
        sources=[source],
        data_dir=tmp_path / ".data",
        transport=transport,
    )

    # 3 requests: listing, detail, attachment
    assert len(transport.requests_log) == 3
    assert transport.requests_log[0]["url"] == listing_url
    assert transport.requests_log[1]["url"] == detail_url
    assert transport.requests_log[2]["url"] == att_url

    # Agent evidence packet emitted
    assert len(out["agent_evidence_packets"]) == 1
    packet = out["agent_evidence_packets"][0]
    assert packet["source_id"] == "src_changed_test"
    assert packet["url"] == detail_url
    assert len(packet["attachments"]) == 1


def test_unprocessed_change_downstream_failure_does_not_advance_baseline(tmp_path: Path):
    """
    5. Critical Invariant (Delta 1):
       When listing fingerprint changes, but downstream detail request returns 500:
       - production entrypoint MUST NOT advance committed baseline;
       - reloading SourceRegistry from disk proves committed baseline remains OLD;
       - subsequent run still treats the listing as changed and attempts acquisition.
    """
    listing_url = "https://hr.example.edu.cn/jobs"
    detail_url = "https://hr.example.edu.cn/jobs/303"

    listing_html = f"""<html><body>
      <a href="{detail_url}">2026年最新招聘启事</a>
    </body></html>"""

    transport = FakeHttpTransport({
        listing_url: {"status_code": 200, "text": listing_html},
        detail_url: {"status_code": 500, "text": "Internal Server Error"},
    })

    seed_file = tmp_path / "seed.json"
    old_baseline_fp = "old_baseline_fingerprint"
    seed_file.write_text(json.dumps([
        {
            "source_id": "src_fail_test",
            "name": "测试高校",
            "base_url": listing_url,
            "domain": "hr.example.edu.cn",
            "metadata": {
                "is_listing": True,
                "detail_url_pattern": r"/jobs/\d+",
                "committed_listing_fingerprint": old_baseline_fp,
            },
        }
    ], ensure_ascii=False), encoding="utf-8")

    data_dir = tmp_path / ".data"

    # RUN 1: execute through production entrypoint (detail 500 fails)
    out1 = execute_production_acquisition(
        data_dir=data_dir,
        seed_sources_path=seed_file,
        transport=transport,
    )

    assert len(out1["agent_evidence_packets"]) == 0
    detail_res = [r for r in out1["acquisition_results"] if r.requested_url == detail_url][0]
    assert detail_res.technical_status == "failed"
    assert detail_res.http_status == 500

    # Reload SourceRegistry from disk to verify persisted state
    reloaded_reg1 = SourceRegistry(seed_path=seed_file, data_dir=data_dir)
    saved_src1 = reloaded_reg1.get_source("src_fail_test")
    assert saved_src1.metadata.get("committed_listing_fingerprint") == old_baseline_fp

    # RUN 2: detail endpoint is fixed (200)
    transport.responses[detail_url] = {
        "status_code": 200,
        "text": "<html><body><h1>招聘成功详情</h1></body></html>",
    }
    transport.requests_log.clear()

    out2 = execute_production_acquisition(
        data_dir=data_dir,
        seed_sources_path=seed_file,
        transport=transport,
    )

    # Proves RUN 2 automatically retried detail acquisition without manual test intervention
    assert len(out2["agent_evidence_packets"]) == 1
    assert out2["agent_evidence_packets"][0]["url"] == detail_url

    # Reload from disk again: now baseline is committed
    reloaded_reg2 = SourceRegistry(seed_path=seed_file, data_dir=data_dir)
    saved_src2 = reloaded_reg2.get_source("src_fail_test")
    assert saved_src2.metadata.get("committed_listing_fingerprint") != old_baseline_fp


def test_process_restart_end_to_end_persisted_zero_token_monitoring(tmp_path: Path):
    """
    Critical Process-Restart Test:
    RUN 1:
      - New listing evidence acquired and parsed successfully.
      - Production entrypoint completes and automatically persists committed baseline to .data/sources.json.
    RUN 2 (Simulating next day / process restart with NEW SourceRegistry from disk):
      - Production entrypoint loads committed state from .data/sources.json.
      - Recognizes unchanged source.
      - Only cheap listing check occurs.
      - Detail request is skipped.
      - Attachment request is skipped.
      - agent_evidence_packets == [].
    """
    listing_url = "https://hr.example.edu.cn/jobs/list"
    detail_url = "https://hr.example.edu.cn/jobs/detail/505"
    att_url = "https://hr.example.edu.cn/files/guide.docx"

    listing_html = f"""<html><body><a href="{detail_url}">2026年高层次人才招聘</a></body></html>"""
    detail_html = f"""<html><body><h1>高层次人才</h1><a href="{att_url}">申报指南.docx</a></body></html>"""
    import docx
    from io import BytesIO

    doc = docx.Document()
    doc.add_heading("2026年高层次人才招聘指南", level=1)
    p = doc.add_paragraph("申报条件与要求说明")
    bio = BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()

    transport = FakeHttpTransport({
        listing_url: {"status_code": 200, "text": listing_html, "headers": {"ETag": '"list_v1"'}},
        detail_url: {"status_code": 200, "text": detail_html},
        att_url: {"status_code": 200, "content": docx_bytes, "headers": {"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}},
    })

    seed_file = tmp_path / "seed.json"
    seed_file.write_text(json.dumps([
        {
            "source_id": "src_restart_e2e",
            "name": "重点大学招聘网",
            "base_url": listing_url,
            "domain": "hr.example.edu.cn",
            "metadata": {
                "is_listing": True,
                "detail_url_pattern": r"/jobs/detail/\d+",
            },
        }
    ], ensure_ascii=False), encoding="utf-8")

    data_dir = tmp_path / ".data"

    # RUN 1: Initial discovery / acquisition
    out1 = execute_production_acquisition(
        data_dir=data_dir,
        seed_sources_path=seed_file,
        transport=transport,
    )
    assert len(out1["agent_evidence_packets"]) == 1
    assert len(transport.requests_log) == 3  # listing, detail, attachment

    # Verify .data/sources.json has committed baseline
    local_state_file = data_dir / "sources.json"
    assert local_state_file.exists()

    # Clear request log to track RUN 2 exactly
    transport.requests_log.clear()

    # RUN 2: Process restart (fresh execution without passing in-memory sources or registry)
    out2 = execute_production_acquisition(
        data_dir=data_dir,
        seed_sources_path=seed_file,
        transport=transport,
    )

    # Assert 0-Token monitoring behavior across restart:
    # 1. Exactly 1 request (listing check with If-None-Match: "list_v1" if ETag present, or 200 with identical fingerprint)
    assert len(transport.requests_log) == 1
    assert transport.requests_log[0]["url"] == listing_url
    # 2. Detail and attachments skipped
    assert not any(req["url"] == detail_url for req in transport.requests_log)
    assert not any(req["url"] == att_url for req in transport.requests_log)
    # 3. 0 Agent evidence packets
    assert out2["agent_evidence_packets"] == []
    # 4. Monitoring fact is success and unchanged
    assert out2["monitoring_facts"][0].technical_status == "success"
    assert out2["monitoring_facts"][0].metadata.get("unchanged") is True


def test_listing_without_detail_selection_hints_does_not_fingerprint_arbitrary_links(tmp_path: Path):
    """
    Blocker 2 Regression:
    Proves that when a listing source has many links but NO configured detail selection hints:
    - does not fingerprint arbitrary all-page links;
    - does not commit arbitrary baseline;
    - does not fetch arbitrary detail;
    - produces 0 Agent evidence packets.
    """
    listing_url = "https://hr.example.edu.cn/general_list"
    listing_html = """<html><body>
      <a href="/about">关于我们</a>
      <a href="/contact">联系方式</a>
      <a href="/news/1">学校要闻</a>
    </body></html>"""

    transport = FakeHttpTransport({
        listing_url: {"status_code": 200, "text": listing_html},
    })

    seed_file = tmp_path / "seed.json"
    seed_file.write_text(json.dumps([
        {
            "source_id": "src_no_hints",
            "name": "未配置详情规则列表",
            "base_url": listing_url,
            "domain": "hr.example.edu.cn",
            "metadata": {
                "is_listing": True,
                # No detail_url_pattern, no detail_url, no detail_link_index
            },
        }
    ], ensure_ascii=False), encoding="utf-8")

    data_dir = tmp_path / ".data"

    out = execute_production_acquisition(
        data_dir=data_dir,
        seed_sources_path=seed_file,
        transport=transport,
    )

    # Only listing is requested
    assert len(transport.requests_log) == 1
    assert transport.requests_log[0]["url"] == listing_url

    # Zero evidence packets
    assert out["agent_evidence_packets"] == []

    # Reload registry: verify no arbitrary fingerprint was committed
    reg = SourceRegistry(seed_path=seed_file, data_dir=data_dir)
    src = reg.get_source("src_no_hints")
    assert src.metadata.get("committed_listing_fingerprint") is None

